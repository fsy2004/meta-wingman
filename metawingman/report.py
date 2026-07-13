# -*- coding: utf-8 -*-
"""一键报告:用 python-docx 从一次运行的产物(图+表)生成 Word 报告——
标题 + 方法(软件/包/参数/复现命令)+ 结果(嵌图+表)+ 参考文献(R citation() 生成,不臆造)。
覆盖全部方法(只用已产出的 png/csv);无需 pandoc。"""
from __future__ import annotations
import csv
import os
import subprocess

from .rlocate import find_rscript

CREATE_NO_WINDOW = 0x08000000 if os.name == "nt" else 0

# 按分类挑该分析实际用到的核心 R 包,交给 R citation() 出真实引用
_GROUP_PKGS = {
    "synthesis": ["metafor", "meta"], "heterogeneity": ["metafor", "meta"],
    "bias": ["metafor", "meta", "metasens"], "robustness": ["metafor"],
    "complex": ["metafor", "dosresmeta", "robumeta", "clubSandwich"],
    "nma": ["netmeta"], "diagnostics": ["mada"], "sequential": ["RTSA"],
    "certainty": ["EValue", "meta"], "reporting": ["meta"], "data": ["estmeansd", "metafor"],
}


def _citations(packages):
    rs = find_rscript()
    if not rs or not packages:
        return []
    pk = ",".join('"%s"' % p for p in packages)
    code = ('for (p in c(%s)) if (requireNamespace(p, quietly=TRUE)) '
            'cat("- ", format(citation(p)[1], style="text"), "\\n", sep="")') % pk
    try:
        out = subprocess.run([rs, "-e", code], capture_output=True, encoding="utf-8",
                             errors="replace", timeout=60, creationflags=CREATE_NO_WINDOW)
        return [l.strip() for l in (out.stdout or "").splitlines() if l.strip().startswith("-")]
    except Exception:
        return []


def _r_version():
    rs = find_rscript()
    try:
        out = subprocess.run([rs, "-e", "cat(R.version.string)"], capture_output=True,
                             encoding="utf-8", errors="replace", timeout=30, creationflags=CREATE_NO_WINDOW)
        return (out.stdout or "R").strip()
    except Exception:
        return "R"


def _read_command(outdir):
    p = os.path.join(outdir, "reproduce.R")
    try:
        for line in open(p, encoding="utf-8"):
            if line.startswith("## 生成:") and "分析:" in line:
                return line.split("分析:", 1)[1].strip()
    except Exception:
        pass
    return ""


def build_report(manifest, params, outputs, outdir, dst):
    from docx import Document
    from docx.shared import Inches, Pt

    title_en = manifest.get("title_en") or manifest.get("id")
    title_zh = manifest.get("title") or ""
    doc = Document()
    doc.add_heading("%s" % title_en, level=1)
    if title_zh and title_zh != title_en:
        doc.add_paragraph(title_zh)

    doc.add_heading("Methods", level=2)
    pkgs = _GROUP_PKGS.get(manifest.get("group", ""), ["metafor", "meta"])
    m = doc.add_paragraph()
    m.add_run("Software. ").bold = True
    m.add_run("Analyses were performed in %s using the %s package(s)."
              % (_r_version(), ", ".join(pkgs)))
    if params:
        pstr = "; ".join("%s = %s" % (k, v) for k, v in params.items() if str(v) != "")
        if pstr:
            p = doc.add_paragraph()
            p.add_run("Settings. ").bold = True
            p.add_run(pstr)
    cmd = _read_command(outdir)
    if cmd:
        p = doc.add_paragraph()
        p.add_run("Reproducibility. ").bold = True
        p.add_run("The exact call is saved in reproduce.R (data as data.csv): ")
        p.add_run(cmd).font.size = Pt(8)

    doc.add_heading("Results", level=2)
    imgs = [o for o in outputs if o.lower().endswith(".png")]
    tbls = [o for o in outputs if o.lower().endswith(".csv") and os.path.basename(o) != "data.csv"]
    for png in imgs:
        try:
            doc.add_picture(png, width=Inches(6.0))
            cap = doc.add_paragraph(os.path.basename(png))
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
        except Exception:
            pass
    for c in tbls:
        try:
            with open(c, encoding="utf-8-sig", newline="") as f:
                rows = list(csv.reader(f))
            if not rows:
                continue
            doc.add_paragraph(os.path.basename(c)).runs[0].italic = True
            t = doc.add_table(rows=1, cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for j, h in enumerate(rows[0]):
                t.rows[0].cells[j].text = str(h)
            for r in rows[1:200]:
                cells = t.add_row().cells
                for j, v in enumerate(r[:len(rows[0])]):
                    cells[j].text = str(v)
        except Exception:
            pass

    cites = _citations(pkgs)
    if cites:
        doc.add_heading("References", level=2)
        for c in cites:
            doc.add_paragraph(c[2:] if c.startswith("- ") else c)

    doc.save(dst)
    return dst
